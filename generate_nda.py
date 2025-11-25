#!/usr/bin/env python3
"""
Script to generate CultivateDynamics_RiskLo_NDA.docx
Requires: pip install python-docx
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx library not installed.")
    print("Please install it with: pip install python-docx")
    exit(1)

# Create document
doc = Document()

# Set margins to 1 inch
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('Mutual Non-Disclosure Agreement (NDA)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(16)
title_run.font.name = 'Times New Roman'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# Opening clause
p = doc.add_paragraph()
p.add_run('This Mutual Non-Disclosure Agreement ("Agreement") is entered into as of [Date], by and between Cultivate Dynamics LLC, the owner and developer of the software product known as \'RiskLo,\' with an address at [Your Address] ("Disclosing Party"), and [Friend\'s Name], with an address at [Their Address] ("Receiving Party").')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Section 1
doc.add_heading('1. Definition of Confidential Information', level=2)
p = doc.add_paragraph()
p.add_run('"Confidential Information" includes proprietary information related to the software product \'RiskLo,\' including but not limited to: business logic, algorithms, computational methods, UI/UX designs, architecture, data processing techniques, business model, pricing structure, user feedback, roadmaps, marketing strategies, and technical implementation details. Confidential Information also includes any information or material disclosed by the Disclosing Party to the Receiving Party that is designated as confidential or that reasonably should be understood to be confidential.')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Section 2
doc.add_heading('2. Obligations of the Receiving Party', level=2)
p = doc.add_paragraph()
p.add_run('The Receiving Party agrees:')
p_format = p.paragraph_format
p_format.space_after = Pt(12)

p = doc.add_paragraph()
p.add_run('(a) to hold and maintain the Confidential Information in strict confidence for the sole benefit of the Disclosing Party;')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.left_indent = Inches(0.25)

p = doc.add_paragraph()
p.add_run('(b) not to disclose the Confidential Information to any third party without the Disclosing Party\'s prior written consent;')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.left_indent = Inches(0.25)

p = doc.add_paragraph()
p.add_run('(c) not to use any of the Confidential Information for any purpose other than privately reviewing or testing RiskLo; and')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.left_indent = Inches(0.25)

p = doc.add_paragraph()
p.add_run('(d) not to copy, store, duplicate, or archive any interface, code, visual layout, or logic from RiskLo without written permission.')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.left_indent = Inches(0.25)

# Section 3
doc.add_heading('3. Exclusions from Confidential Information', level=2)
p = doc.add_paragraph()
p.add_run('Confidential Information does not include information that:')
p_format = p.paragraph_format
p_format.space_after = Pt(12)

p = doc.add_paragraph()
p.add_run('(a) is already publicly available through no fault of the Receiving Party;')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.left_indent = Inches(0.25)

p = doc.add_paragraph()
p.add_run('(b) is received from a third party without restriction or breach of a confidentiality obligation; or')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.left_indent = Inches(0.25)

p = doc.add_paragraph()
p.add_run('(c) is independently developed by the Receiving Party without use of or reference to the Disclosing Party\'s Confidential Information.')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.left_indent = Inches(0.25)

# Section 4
doc.add_heading('4. Duration', level=2)
p = doc.add_paragraph()
p.add_run('The obligations of this Agreement will survive for a period of two (2) years from the date of disclosure of the Confidential Information.')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Section 5
doc.add_heading('5. Ownership', level=2)
p = doc.add_paragraph()
p.add_run('The Receiving Party acknowledges that all intellectual property rights, ownership rights, and commercial rights to RiskLo belong exclusively to Cultivate Dynamics LLC. No license or ownership interest is transferred or granted under this Agreement.')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Section 6
doc.add_heading('6. Non-Compete / No Reverse Engineering', level=2)
p = doc.add_paragraph()
p.add_run('The Receiving Party agrees not to reverse engineer, decompile, disassemble, analyze, or derive source code, computational mechanisms, or algorithms from RiskLo.')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph()
p.add_run('The Receiving Party further agrees not to develop, market, sell, or distribute a competing software product or service that performs substantially similar functions to RiskLo or targets an overlapping user base for a period of three (3) years.')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Section 7
doc.add_heading('7. Electronic Signature Consent', level=2)
p = doc.add_paragraph()
p.add_run('Both parties agree that this Agreement may be executed electronically, and that digital signatures have the same legal effect as handwritten signatures.')
p_format = p.paragraph_format
p_format.space_after = Pt(12)
p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Signature sections
doc.add_paragraph()  # Spacing

# Disclosing Party signature block
p = doc.add_paragraph()
p.add_run('Cultivate Dynamics LLC (Disclosing Party)').bold = True
p_format = p.paragraph_format
p_format.space_before = Pt(36)
p_format.space_after = Pt(6)

p = doc.add_paragraph()
p.add_run('Representing the product: RiskLo')
p_format = p.paragraph_format
p_format.space_after = Pt(12)

p = doc.add_paragraph()
p.add_run('Signed electronically by:')
p_format = p.paragraph_format
p_format.space_after = Pt(6)

# Signature line
p = doc.add_paragraph()
p_format = p.paragraph_format
p_format.space_after = Pt(12)
run = p.add_run('_' * 50)
run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('Date:')
p_format = p.paragraph_format
p_format.space_after = Pt(6)

# Date line
p = doc.add_paragraph()
p_format = p.paragraph_format
p_format.space_after = Pt(12)
run = p.add_run('_' * 25)
run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('IP Address (optional):')
p_format = p.paragraph_format
p_format.space_after = Pt(6)

# IP line
p = doc.add_paragraph()
p_format = p.paragraph_format
p_format.space_after = Pt(24)
run = p.add_run('_' * 30)
run.font.size = Pt(12)

# Receiving Party signature block
p = doc.add_paragraph()
p.add_run('Receiving Party').bold = True
p_format = p.paragraph_format
p_format.space_before = Pt(36)
p_format.space_after = Pt(12)

p = doc.add_paragraph()
p.add_run('Signed electronically by:')
p_format = p.paragraph_format
p_format.space_after = Pt(6)

# Signature line
p = doc.add_paragraph()
p_format = p.paragraph_format
p_format.space_after = Pt(12)
run = p.add_run('_' * 50)
run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('Date:')
p_format = p.paragraph_format
p_format.space_after = Pt(6)

# Date line
p = doc.add_paragraph()
p_format = p.paragraph_format
p_format.space_after = Pt(12)
run = p.add_run('_' * 25)
run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('IP Address (optional):')
p_format = p.paragraph_format
p_format.space_after = Pt(6)

# IP line
p = doc.add_paragraph()
p_format = p.paragraph_format
p_format.space_after = Pt(12)
run = p.add_run('_' * 30)
run.font.size = Pt(12)

# Set default font to Times New Roman
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set heading styles
for heading_style in ['Heading 1', 'Heading 2']:
    try:
        style = doc.styles[heading_style]
        font = style.font
        font.name = 'Times New Roman'
    except:
        pass

# Save document
output_file = 'CultivateDynamics_RiskLo_NDA.docx'
doc.save(output_file)
print(f"Document created successfully: {output_file}")

